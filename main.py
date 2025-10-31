import os
import re
import shutil
import time
import threading
import subprocess
from tkinterdnd2 import TkinterDnD, DND_FILES
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, ttk
from pathlib import Path
from docx2pdf import convert as convert_docx_to_pdf
from docx import Document
from pptx import Presentation
from comtypes.client import CreateObject
from PyPDF2 import PdfMerger
from PIL import Image

# =================== Utility Functions ===================

def extract_number(filename):
    """Extract a number from the filename (used for sorting)."""
    match = re.search(r'(\d+)', filename)
    return int(match.group(1)) if match else float('inf')

def convert_docx_to_txt(docx_path, txt_path):
    """Convert DOCX to plain TXT."""
    doc = Document(docx_path)
    with open(txt_path, "w", encoding="utf-8") as f:
        for para in doc.paragraphs:
            f.write(para.text + "\n")

def convert_docx_to_html(docx_path, html_path):
    """Convert DOCX to basic HTML format."""
    doc = Document(docx_path)
    with open(html_path, "w", encoding="utf-8") as f:
        f.write("<html><body>\n")
        for para in doc.paragraphs:
            f.write(f"<p>{para.text}</p>\n")
        f.write("</body></html>")

def convert_pptx_to_pdf(pptx_path, pdf_path):
    """Convert PPTX to PDF using PowerPoint COM interface (Windows only)."""
    powerpoint = CreateObject("Powerpoint.Application")
    powerpoint.Visible = 1
    try:
        deck = powerpoint.Presentations.Open(str(pptx_path), WithWindow=False)
        deck.SaveAs(str(pdf_path), 32)  # 32 = PDF format
        deck.Close()
    finally:
        powerpoint.Quit()

def open_folder(path):
    """Open a folder in the system file explorer."""
    if os.name == 'nt':
        os.startfile(str(path))
    elif os.name == 'posix':
        subprocess.run(["xdg-open", str(path)])

# =================== Text Resources ===================

texts = {
    "select_files": "Select files to convert",
    "select_folder": "Select folder to convert",
    "select_output": "Select output folder",
    "converting": "Converting... please wait.",
    "done": "✅ Conversion complete!",
    "fill_fields": "Please fill all fields before converting.",
    "error": "Missing Info",
    "merge_title": "Merged PDF",
    "merge_prompt": "Enter name for merged PDF (without extension):",
    "cancel": "Canceled",
    "no_merge_name": "No name provided for merged PDF. Canceling."
}

selection_listbox = None


def refresh_selection_display(items, mode="files"):
    """Update the on-screen list of selected items."""
    if selection_listbox is None:
        return

    selection_listbox.delete(0, tk.END)

    if not items:
        selection_listbox.insert(tk.END, "No items selected yet.")
        return

    if mode == "folder":
        folder_path = Path(items[0])
        selection_listbox.insert(tk.END, f"[Folder] {folder_path}")
        try:
            children = sorted(
                list(folder_path.iterdir()),
                key=lambda p: (p.is_file(), p.name.lower())
            )
            selection_listbox.insert(tk.END, f"  {len(children)} item(s) inside")
        except Exception:
            selection_listbox.insert(tk.END, "  (Unable to preview folder contents)")
            return

        max_items = 25
        for idx, child in enumerate(children):
            prefix = "  • " if child.is_file() else "  📁 "
            selection_listbox.insert(tk.END, f"{prefix}{child.name}")
            if idx + 1 == max_items:
                selection_listbox.insert(tk.END, "  ... (more items not shown)")
                break
    else:
        selection_listbox.insert(tk.END, f"{len(items)} item(s) selected:")
        for path in items:
            p = Path(path)
            label = f"  📁 {p.name}" if p.is_dir() else f"  • {p.name}"
            selection_listbox.insert(tk.END, label)


def clear_selection():
    """Clear the current input selection and reset the list display."""
    input_var.set("")
    refresh_selection_display([])

# =================== GUI Actions ===================

def browse_input():
    """Open file/folder selector depending on mode."""
    if input_choice.get() == "files":
        files = filedialog.askopenfilenames(title=texts["select_files"])
        input_var.set("\n".join(files))
        if files:
            refresh_selection_display(files, mode="files")
        else:
            refresh_selection_display([])
    else:
        folder = filedialog.askdirectory(title=texts["select_folder"])
        input_var.set(folder)
        if folder:
            refresh_selection_display([folder], mode="folder")
        else:
            refresh_selection_display([])

def browse_output():
    """Open folder selector for output location."""
    folder = filedialog.askdirectory(title=texts["select_output"])
    output_var.set(folder)

def handle_drop(event):
    """Handle drag & drop file or folder input."""
    paths = root.tk.splitlist(event.data)
    clean = [Path(p.strip('{}')) for p in paths if Path(p.strip('{}')).exists()]
    if clean:
        if clean[0].is_dir():
            input_choice.set("folder")
            folder_path = str(clean[0])
            input_var.set(folder_path)
            refresh_selection_display([folder_path], mode="folder")
        else:
            input_choice.set("files")
            files = [str(p) for p in clean]
            input_var.set("\n".join(files))
            refresh_selection_display(files, mode="files")

def threaded_conversion():
    """Prepare GUI interaction before background conversion."""
    if merge_var.get() and format_vars["pdf"].get():
        merged_name = simpledialog.askstring(texts["merge_title"], texts["merge_prompt"])
        if not merged_name:
            messagebox.showwarning(texts["cancel"], texts["no_merge_name"])
            return
    else:
        merged_name = None

    threading.Thread(target=start_conversion, args=(merged_name,), daemon=True).start()

def start_conversion(merged_name=None):
    """Run the actual file conversion."""
    log = []
    progress['value'] = 0
    progress_label.config(text=texts["converting"])
    root.update_idletasks()

    inputs = input_var.get().strip()
    output = Path(output_var.get().strip())
    selected_formats = [fmt for fmt, var in format_vars.items() if var.get()]
    merge = merge_var.get()

    if not inputs or not output or not selected_formats:
        messagebox.showerror(texts["error"], texts["fill_fields"])
        return

    input_paths = []
    if '\n' in inputs:
        input_paths = [Path(p) for p in inputs.split('\n') if Path(p).exists()]
    else:
        p = Path(inputs)
        input_paths = list(p.glob("*")) if p.is_dir() else [p]

    image_types = {"jpg", "jpeg", "png", "bmp", "gif", "tiff"}
    pdfs_created = []

    total_steps = len(input_paths) * len(selected_formats)
    progress['maximum'] = total_steps
    step = 0

    for file in input_paths:
        ext = file.suffix.lower().lstrip('.')
        stem = file.stem


        for fmt in selected_formats:
            out_dir = output / f"{fmt}_output"
            out_dir.mkdir(exist_ok=True)
            save_path = out_dir / f"{stem}.{fmt}"
            counter = 1
            while save_path.exists():
                save_path = out_dir / f"{stem}_{counter}.{fmt}"
                counter += 1

            try:
                if ext == "docx":
                    if fmt == "pdf":
                        convert_docx_to_pdf(str(file), str(save_path))
                        pdfs_created.append(save_path)
                    elif fmt == "txt":
                        convert_docx_to_txt(file, save_path)
                    elif fmt == "html":
                        convert_docx_to_html(file, save_path)

                # Add support for PPT to PDF
                elif (ext == "pptx" or ext == "ppt") and fmt == "pdf":
                    convert_pptx_to_pdf(file, save_path)
                    time.sleep(1)
                    pdfs_created.append(save_path)

                elif ext == "pdf" and fmt == "pdf":
                    shutil.copy(file, save_path)
                    pdfs_created.append(save_path)

                elif ext in image_types and fmt in image_types and fmt != ext:
                    with Image.open(file) as img:
                        rgb_img = img.convert("RGB") if img.mode in ("RGBA", "P") else img
                        rgb_img.save(save_path, fmt.upper())

                log.append(f"[OK] {file.name} -> {save_path.name}")
            except Exception as e:
                log.append(f"[FAIL] {file.name} to {fmt.upper()} – {e}")

            step += 1
            progress['value'] = step
            root.update_idletasks()

    if merge and "pdf" in selected_formats:
        try:
            merger = PdfMerger()
            for pdf in sorted(pdfs_created, key=lambda f: extract_number(f.name)):
                merger.append(str(pdf))
            merged_path = output / f"{merged_name}.pdf"
            counter = 1
            while merged_path.exists():
                merged_path = output / f"{merged_name}_{counter}.pdf"
                counter += 1
            merger.write(str(merged_path))
            merger.close()
            log.append(f"[MERGED] PDF saved to: {merged_path}")
        except Exception as e:
            log.append(f"[ERROR] Merging failed: {e}")

    with open(output / "conversion_log.txt", "w", encoding="utf-8") as f:
        f.write("\n".join(log))

    progress_label.config(text=texts["done"])
    open_folder(output)

# =================== GUI Setup ===================

root = TkinterDnD.Tk()
root.title("🛠️ File Converter Tool")
root.geometry("720x900")
root.minsize(620, 720)

input_choice = tk.StringVar(value="files")
input_var = tk.StringVar()
output_var = tk.StringVar()
merge_var = tk.BooleanVar()
format_vars = {fmt: tk.BooleanVar() for fmt in ["pdf", "txt", "html", "jpg", "png", "bmp", "tiff"]}

main_frame = tk.Frame(root, padx=20, pady=20)
main_frame.pack(fill="both", expand=True)

header_label = tk.Label(main_frame, text="File Converter Tool", font=("Segoe UI", 18, "bold"))
header_label.pack(anchor="center")
subheader_label = tk.Label(
    main_frame,
    text="Convert documents, presentations, images, and more with a single click.",
    font=("Segoe UI", 10)
)
subheader_label.pack(anchor="center", pady=(4, 18))

input_frame = tk.LabelFrame(main_frame, text="Input selection", padx=12, pady=12)
input_frame.pack(fill="x", pady=(0, 15))

choice_frame = tk.Frame(input_frame)
choice_frame.pack(anchor="w", pady=(0, 8))
tk.Radiobutton(choice_frame, text="Select Files", variable=input_choice, value="files").pack(side="left", padx=(0, 15))
tk.Radiobutton(choice_frame, text="Select Folder", variable=input_choice, value="folder").pack(side="left")

controls_frame = tk.Frame(input_frame)
controls_frame.pack(fill="x")
tk.Button(controls_frame, text="📂 Browse Input", command=browse_input).pack(side="left")
tk.Button(controls_frame, text="Clear selection", command=clear_selection).pack(side="right")

tk.Entry(input_frame, textvariable=input_var, width=70).pack(fill="x", pady=(10, 0))

selection_frame = tk.LabelFrame(main_frame, text="Selected items", padx=12, pady=12)
selection_frame.pack(fill="both", expand=False, pady=(0, 15))
selection_container = tk.Frame(selection_frame)
selection_container.pack(fill="both", expand=True)
selection_listbox = tk.Listbox(selection_container, height=10, activestyle="none")
selection_listbox.pack(side="left", fill="both", expand=True)
selection_scrollbar = tk.Scrollbar(selection_container, orient="vertical", command=selection_listbox.yview)
selection_scrollbar.pack(side="right", fill="y")
selection_listbox.config(yscrollcommand=selection_scrollbar.set)

output_frame = tk.LabelFrame(main_frame, text="Output folder", padx=12, pady=12)
output_frame.pack(fill="x", pady=(0, 15))
tk.Button(output_frame, text="📁 Browse Output", command=browse_output).pack(anchor="w")
tk.Entry(output_frame, textvariable=output_var, width=70).pack(fill="x", pady=(8, 0))

formats_frame = tk.LabelFrame(main_frame, text="Select output formats", padx=12, pady=12)
formats_frame.pack(fill="x", pady=(0, 15))
formats_grid = tk.Frame(formats_frame)
formats_grid.pack(anchor="w")
for idx, (fmt, var) in enumerate(format_vars.items()):
    row, col = divmod(idx, 3)
    tk.Checkbutton(formats_grid, text=fmt.upper(), variable=var).grid(row=row, column=col, sticky="w", padx=12, pady=6)

options_frame = tk.LabelFrame(main_frame, text="Options", padx=12, pady=12)
options_frame.pack(fill="x", pady=(0, 15))
tk.Checkbutton(options_frame, text="🗃️ Merge PDFs into one", variable=merge_var).pack(anchor="w")
tk.Button(options_frame, text="🚀 Start Conversion", command=threaded_conversion).pack(fill="x", pady=(10, 0))

progress_frame = tk.LabelFrame(main_frame, text="Progress", padx=12, pady=12)
progress_frame.pack(fill="x")
progress = ttk.Progressbar(progress_frame, mode='determinate')
progress.pack(fill="x", pady=(0, 6))
progress_label = tk.Label(progress_frame, text="Waiting for selection...")
progress_label.pack(anchor="w")

footer_label = tk.Label(main_frame, text="v2.1 | Created by Yoavs5771", font=("Arial", 8))
footer_label.pack(side="bottom", pady=(18, 0))

refresh_selection_display([])

# Enable drag & drop
root.drop_target_register(DND_FILES)
root.dnd_bind('<<Drop>>', handle_drop)

root.mainloop()
